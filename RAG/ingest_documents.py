#!/usr/bin/env python3
"""
Document Ingestion Script for Memvid RAG

This script ingests documents (markdown, text, PDF) into a Memvid memory file
for later retrieval by the RAG system.

Usage:
    python ingest_documents.py --input ./knowledge --output ./memories/infrastructure.mv2
    python ingest_documents.py --input ./docs/guide.pdf --output ./memories/guide.mv2

Author: Oktay Savdi
Date: January 2026
"""

import os
import argparse
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from tqdm import tqdm
from dotenv import load_dotenv

load_dotenv()


@dataclass
class Document:
    """Represents a document to be ingested"""
    content: str
    title: str
    uri: str
    source_file: str
    tags: Dict[str, str]


class DocumentLoader:
    """Load documents from various file formats"""
    
    SUPPORTED_EXTENSIONS = {'.md', '.txt', '.pdf', '.docx', '.rst'}
    
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def load_file(self, file_path: Path) -> List[Document]:
        """Load a single file and return list of document chunks"""
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return self._load_pdf(file_path)
        elif suffix == '.docx':
            return self._load_docx(file_path)
        elif suffix in {'.md', '.txt', '.rst'}:
            return self._load_text(file_path)
        else:
            print(f"⚠️  Skipping unsupported file: {file_path}")
            return []
    
    def _load_text(self, file_path: Path) -> List[Document]:
        """Load plain text or markdown file"""
        try:
            content = file_path.read_text(encoding='utf-8')
            chunks = self._chunk_text(content)
            
            documents = []
            for i, chunk in enumerate(chunks):
                documents.append(Document(
                    content=chunk,
                    title=f"{file_path.stem} (Part {i+1}/{len(chunks)})" if len(chunks) > 1 else file_path.stem,
                    uri=f"mv2://docs/{file_path.name}#chunk-{i+1}",
                    source_file=str(file_path),
                    tags={
                        "type": file_path.suffix.lstrip('.'),
                        "chunk": str(i + 1),
                        "total_chunks": str(len(chunks))
                    }
                ))
            
            return documents
        except Exception as e:
            print(f"❌ Error loading {file_path}: {e}")
            return []
    
    def _load_pdf(self, file_path: Path) -> List[Document]:
        """Load PDF file"""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(str(file_path))
            full_text = ""
            
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text:
                    full_text += f"\n\n--- Page {page_num} ---\n\n{text}"
            
            chunks = self._chunk_text(full_text)
            
            documents = []
            for i, chunk in enumerate(chunks):
                documents.append(Document(
                    content=chunk,
                    title=f"{file_path.stem} (Part {i+1}/{len(chunks)})",
                    uri=f"mv2://docs/{file_path.name}#chunk-{i+1}",
                    source_file=str(file_path),
                    tags={
                        "type": "pdf",
                        "chunk": str(i + 1),
                        "total_chunks": str(len(chunks)),
                        "pages": str(len(reader.pages))
                    }
                ))
            
            return documents
        except ImportError:
            print("❌ pypdf not installed. Run: pip install pypdf")
            return []
        except Exception as e:
            print(f"❌ Error loading PDF {file_path}: {e}")
            return []
    
    def _load_docx(self, file_path: Path) -> List[Document]:
        """Load DOCX file"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(str(file_path))
            full_text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            chunks = self._chunk_text(full_text)
            
            documents = []
            for i, chunk in enumerate(chunks):
                documents.append(Document(
                    content=chunk,
                    title=f"{file_path.stem} (Part {i+1}/{len(chunks)})",
                    uri=f"mv2://docs/{file_path.name}#chunk-{i+1}",
                    source_file=str(file_path),
                    tags={
                        "type": "docx",
                        "chunk": str(i + 1),
                        "total_chunks": str(len(chunks))
                    }
                ))
            
            return documents
        except ImportError:
            print("❌ python-docx not installed. Run: pip install python-docx")
            return []
        except Exception as e:
            print(f"❌ Error loading DOCX {file_path}: {e}")
            return []
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]
            
            # Try to break at sentence/paragraph boundary
            if end < len(text):
                # Look for paragraph break first
                last_para = chunk.rfind('\n\n')
                last_period = chunk.rfind('. ')
                last_newline = chunk.rfind('\n')
                
                # Prefer paragraph break, then sentence, then newline
                if last_para > self.chunk_size // 2:
                    break_point = last_para + 2
                elif last_period > self.chunk_size // 2:
                    break_point = last_period + 2
                elif last_newline > self.chunk_size // 2:
                    break_point = last_newline + 1
                else:
                    break_point = self.chunk_size
                
                chunk = text[start:start + break_point]
                end = start + break_point
            
            chunk = chunk.strip()
            if chunk:
                chunks.append(chunk)
            
            start = max(start + 1, end - self.chunk_overlap)
        
        return chunks
    
    def load_directory(self, directory: Path) -> List[Document]:
        """Recursively load all supported documents from a directory"""
        documents = []
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                docs = self.load_file(file_path)
                documents.extend(docs)
        
        return documents


class MemvidIngester:
    """Ingest documents into Memvid memory"""
    
    def __init__(self, memory_file: str):
        self.memory_file = Path(memory_file)
        self.memory_file.parent.mkdir(parents=True, exist_ok=True)
        self._sdk_available = self._check_sdk()
    
    def _check_sdk(self) -> bool:
        """Check if memvid SDK is available"""
        try:
            import memvid_sdk
            return True
        except ImportError:
            return False
    
    def create_memory(self) -> bool:
        """Create a new memory file"""
        if self._sdk_available:
            try:
                import memvid_sdk
                self.mem = memvid_sdk.create(str(self.memory_file))
                return True
            except Exception as e:
                print(f"❌ SDK error: {e}")
                return self._create_via_cli()
        else:
            return self._create_via_cli()
    
    def _create_via_cli(self) -> bool:
        """Create memory using CLI"""
        import subprocess
        try:
            # Remove existing file if present
            if self.memory_file.exists():
                self.memory_file.unlink()
            
            result = subprocess.run(
                ["memvid", "create", str(self.memory_file)],
                capture_output=True, text=True, timeout=30
            )
            return result.returncode == 0
        except FileNotFoundError:
            print("❌ memvid CLI not found. Using file-based fallback.")
            return self._create_fallback()
        except Exception as e:
            print(f"❌ CLI error: {e}")
            return False
    
    def _create_fallback(self) -> bool:
        """Create a simple JSON-based storage as fallback"""
        try:
            # Create a simple JSON file as fallback
            fallback_file = self.memory_file.with_suffix('.json')
            fallback_file.write_text('{"documents": [], "metadata": {"version": "fallback"}}')
            self.memory_file = fallback_file
            print(f"⚠️  Using JSON fallback: {fallback_file}")
            return True
        except Exception as e:
            print(f"❌ Fallback creation failed: {e}")
            return False
    
    def ingest(self, documents: List[Document], batch_size: int = 10) -> int:
        """Ingest documents into memory"""
        if not documents:
            print("⚠️  No documents to ingest")
            return 0
        
        print(f"\n📥 Ingesting {len(documents)} document chunks...")
        
        success_count = 0
        
        if self._sdk_available:
            success_count = self._ingest_via_sdk(documents)
        elif self.memory_file.suffix == '.json':
            success_count = self._ingest_via_json(documents)
        else:
            success_count = self._ingest_via_cli(documents, batch_size)
        
        print(f"\n✅ Successfully ingested {success_count}/{len(documents)} chunks")
        return success_count
    
    def _ingest_via_sdk(self, documents: List[Document]) -> int:
        """Ingest using Memvid SDK"""
        try:
            import memvid_sdk
            mem = memvid_sdk.open(str(self.memory_file))
            
            success_count = 0
            for doc in tqdm(documents, desc="Ingesting"):
                try:
                    mem.put(
                        doc.content.encode(),
                        title=doc.title,
                        uri=doc.uri,
                        tags=doc.tags
                    )
                    success_count += 1
                except Exception as e:
                    print(f"⚠️  Failed to add: {doc.title} - {e}")
            
            mem.commit()
            return success_count
        except Exception as e:
            print(f"❌ SDK ingestion failed: {e}")
            return 0
    
    def _ingest_via_cli(self, documents: List[Document], batch_size: int) -> int:
        """Ingest using Memvid CLI"""
        import subprocess
        import tempfile
        import json
        
        success_count = 0
        
        for doc in tqdm(documents, desc="Ingesting"):
            with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
                f.write(doc.content)
                temp_path = f.name
            
            try:
                result = subprocess.run(
                    ["memvid", "put", str(self.memory_file), "--input", temp_path],
                    capture_output=True, text=True, timeout=60
                )
                if result.returncode == 0:
                    success_count += 1
            except Exception as e:
                print(f"⚠️  CLI error for {doc.title}: {e}")
            finally:
                os.unlink(temp_path)
        
        return success_count
    
    def _ingest_via_json(self, documents: List[Document]) -> int:
        """Ingest into JSON fallback storage"""
        import json
        
        try:
            data = json.loads(self.memory_file.read_text())
            
            for doc in tqdm(documents, desc="Ingesting"):
                data["documents"].append({
                    "content": doc.content,
                    "title": doc.title,
                    "uri": doc.uri,
                    "source_file": doc.source_file,
                    "tags": doc.tags
                })
            
            self.memory_file.write_text(json.dumps(data, indent=2))
            return len(documents)
        except Exception as e:
            print(f"❌ JSON ingestion failed: {e}")
            return 0


def main():
    parser = argparse.ArgumentParser(
        description="Ingest documents into Memvid memory for RAG",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Ingest all documents from a directory
  python ingest_documents.py --input ./knowledge --output ./memories/infrastructure.mv2

  # Ingest a single PDF
  python ingest_documents.py --input ./docs/guide.pdf --output ./memories/guide.mv2

  # Custom chunk settings
  python ingest_documents.py --input ./docs --output ./memories/docs.mv2 --chunk-size 1000 --overlap 200
        """
    )
    
    parser.add_argument(
        "--input", "-i", required=True,
        help="Input file or directory to ingest"
    )
    parser.add_argument(
        "--output", "-o", required=True,
        help="Output Memvid memory file (.mv2)"
    )
    parser.add_argument(
        "--chunk-size", type=int, default=800,
        help="Size of text chunks (default: 800)"
    )
    parser.add_argument(
        "--overlap", type=int, default=100,
        help="Overlap between chunks (default: 100)"
    )
    parser.add_argument(
        "--force", "-f", action="store_true",
        help="Overwrite existing memory file"
    )
    
    args = parser.parse_args()
    
    input_path = Path(args.input)
    output_path = Path(args.output)
    
    # Validate input
    if not input_path.exists():
        print(f"❌ Input path does not exist: {input_path}")
        return 1
    
    # Check if output exists
    if output_path.exists() and not args.force:
        print(f"⚠️  Output file exists: {output_path}")
        response = input("Overwrite? [y/N]: ").strip().lower()
        if response != 'y':
            print("Aborted.")
            return 0
    
    print("\n" + "="*60)
    print("📚 Document Ingestion for Memvid RAG")
    print("="*60)
    print(f"📁 Input: {input_path}")
    print(f"💾 Output: {output_path}")
    print(f"📏 Chunk size: {args.chunk_size}, Overlap: {args.overlap}")
    print("="*60 + "\n")
    
    # Load documents
    loader = DocumentLoader(
        chunk_size=args.chunk_size,
        chunk_overlap=args.overlap
    )
    
    if input_path.is_file():
        print(f"📄 Loading file: {input_path}")
        documents = loader.load_file(input_path)
    else:
        print(f"📂 Scanning directory: {input_path}")
        documents = loader.load_directory(input_path)
    
    if not documents:
        print("❌ No documents loaded. Check input path and supported formats.")
        return 1
    
    print(f"\n📊 Loaded {len(documents)} document chunks from {input_path}")
    
    # Create memory and ingest
    ingester = MemvidIngester(str(output_path))
    
    if output_path.exists():
        output_path.unlink()
    
    if not ingester.create_memory():
        print("❌ Failed to create memory file")
        return 1
    
    success = ingester.ingest(documents)
    
    if success > 0:
        print(f"\n✅ Ingestion complete!")
        print(f"📁 Memory file: {output_path}")
        print(f"📊 Documents: {success} chunks indexed")
        return 0
    else:
        print("❌ Ingestion failed")
        return 1


if __name__ == "__main__":
    exit(main())
